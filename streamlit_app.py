# -*- coding: utf-8 -*-
"""
FinancePlus Mandato - Web App Streamlit
Versione Cloud OK: non usa tkinter.

Funzioni principali:
- caricamento visura/report PDF;
- estrazione dati azienda/amministratore;
- correzione manuale dei dati estratti;
- compilazione Mandato_vuoto.docx mantenendo il modello Word;
- download DOCX e, se LibreOffice e' disponibile, anche PDF.
"""

from __future__ import annotations

import io
import json
import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import streamlit as st
from docx import Document

try:
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

APP_NAME = "FinancePlus Mandato Web"
BASE_DIR = Path(__file__).resolve().parent
ASSET_DIR = BASE_DIR / "assets"
DEFAULT_TEMPLATE = ASSET_DIR / "Mandato_vuoto.docx"
LOGO_FILE = ASSET_DIR / "LOGO_FINANCE_2.PNG"


# -----------------------------------------------------------------------------
# Utility testo / importi
# -----------------------------------------------------------------------------


def normalize_space(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\u00a0", " ").replace("\t", " ")
    text = re.sub(r"[ \f\v]+", " ", text)
    text = re.sub(r"\s+\n", "\n", text)
    text = re.sub(r"\n\s+", "\n", text)
    return text.strip()


def clean_inline(value: Any) -> str:
    text = normalize_space(value)
    text = text.strip(" :-|;,.\n\t")
    text = re.sub(r"\s{2,}", " ", text)
    return text


def safe_filename(name: str, fallback: str = "mandato") -> str:
    name = clean_inline(name) or fallback
    name = re.sub(r"[^A-Za-z0-9_. -]+", "_", name)
    name = re.sub(r"\s+", "_", name).strip("._-")
    return name[:120] or fallback


def today_it() -> str:
    return datetime.now().strftime("%d/%m/%Y")


def format_euro(raw: str) -> str:
    value = clean_inline(raw)
    if not value:
        return ""
    if "€" in value:
        return value
    return f"€ {value}"


# -----------------------------------------------------------------------------
# Estrazione PDF
# -----------------------------------------------------------------------------


def extract_pdf_text_from_bytes(pdf_bytes: bytes) -> str:
    chunks: List[str] = []

    if pdfplumber is not None:
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        chunks.append(text)
        except Exception:
            pass

    if not chunks and PdfReader is not None:
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            for page in reader.pages:
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append(text)
        except Exception:
            pass

    return normalize_space("\n".join(chunks))


def lines_from_text(text: str) -> List[str]:
    return [clean_inline(x) for x in text.splitlines() if clean_inline(x)]


def first_match(patterns: Iterable[str], text: str, flags: int = re.I | re.M) -> str:
    for pattern in patterns:
        match = re.search(pattern, text, flags)
        if match:
            return clean_inline(match.group(1))
    return ""


def find_near_label(lines: List[str], labels: Iterable[str], max_lines: int = 3) -> str:
    labels_l = [x.lower() for x in labels]
    for i, line in enumerate(lines):
        low = line.lower()
        for label in labels_l:
            if label in low:
                after = re.split(re.escape(label), line, flags=re.I)[-1]
                after = clean_inline(after)
                after = re.sub(r"^[:\-–]+", "", after).strip()
                if after and after.lower() != label:
                    return after
                values = []
                for j in range(i + 1, min(len(lines), i + 1 + max_lines)):
                    nxt = lines[j]
                    if re.search(r"^(partita iva|codice fiscale|pec|sede|rea|telefono|email|attivita|oggetto sociale)\b", nxt, re.I):
                        break
                    values.append(nxt)
                return clean_inline(" ".join(values))
    return ""


def extract_denominazione(lines: List[str], text: str) -> str:
    value = first_match([
        r"(?:Denominazione|Ragione sociale|Impresa)\s*[:\-]?\s*([^\n]{3,120})",
        r"Dati identificativi\s+([^\n]{3,120})",
    ], text)
    value = re.sub(r"\s+(Partita IVA|Codice fiscale|REA)\b.*$", "", value, flags=re.I).strip()
    if value:
        return value.upper()

    for line in lines[:80]:
        if re.search(r"\b(S\.?R\.?L\.?|SPA|S\.?P\.?A\.?|SOCIETA'|SOCIETÀ|SNC|SAS)\b", line, re.I):
            if not re.search(r"registro|camera|visura|rea|partita|codice fiscale", line, re.I):
                return line.upper()
    return ""


def extract_piva(text: str) -> str:
    value = first_match([
        r"Partita\s+IVA\s*[:\-]?\s*([0-9]{11})",
        r"P\.?\s*IVA\s*[:\-]?\s*([0-9]{11})",
        r"IVA\s*[:\-]?\s*([0-9]{11})",
    ], text)
    if value:
        return re.sub(r"\D+", "", value)
    all_codes = re.findall(r"(?<!\d)(\d{11})(?!\d)", text)
    return all_codes[0] if all_codes else ""


def extract_company_cf(text: str, piva: str = "") -> str:
    value = first_match([
        r"Codice\s+fiscale\s*[:\-]?\s*([0-9]{11})",
        r"C\.?F\.?\s*[:\-]?\s*([0-9]{11})",
    ], text)
    if value:
        return re.sub(r"\D+", "", value)
    return piva


def extract_emails(text: str) -> Tuple[str, str]:
    emails = re.findall(r"[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}", text, flags=re.I)
    cleaned: List[str] = []
    for email in emails:
        e = email.strip(".,;:").lower()
        if e not in cleaned:
            cleaned.append(e)
    pec = ""
    mail = ""
    for email in cleaned:
        if "pec" in email or "legalmail" in email or "arubapec" in email:
            pec = pec or email
        else:
            mail = mail or email
    return mail, pec


def extract_sede_legale(lines: List[str], text: str) -> str:
    value = first_match([
        r"Sede\s+legale\s*[:\-]?\s*([^\n]{5,180})",
        r"Indirizzo\s+sede\s+legale\s*[:\-]?\s*([^\n]{5,180})",
    ], text)
    if value:
        return value
    return find_near_label(lines, ["sede legale", "indirizzo sede legale"], max_lines=2)


def extract_ateco(lines: List[str], text: str) -> Tuple[str, str]:
    code = first_match([
        r"(?:ATECO|Codice attivita')\s*[:\-]?\s*([0-9]{2}(?:\.[0-9]{1,2}){0,3})",
        r"Codice\s+attivit[aà]\s*[:\-]?\s*([0-9]{2}(?:\.[0-9]{1,2}){0,3})",
    ], text)
    description = ""
    for i, line in enumerate(lines):
        if re.search(r"ateco|attivit[aà]\s+prevalente|codice attivit", line, re.I):
            piece = re.sub(r".*?(ateco|attivit[aà]\s+prevalente|codice attivit[aà])\s*[:\-]?", "", line, flags=re.I).strip()
            if piece and not code:
                code_match = re.search(r"([0-9]{2}(?:\.[0-9]{1,2}){0,3})", piece)
                if code_match:
                    code = code_match.group(1)
            if i + 1 < len(lines):
                description = clean_inline(lines[i + 1])
            break
    return code, description


def extract_admin(lines: List[str], text: str) -> Dict[str, str]:
    admin_block = "\n".join(lines)
    idx = 0
    for i, line in enumerate(lines):
        if re.search(r"amministratore|legale rappresentante|titolare|consiglio di amministrazione", line, re.I):
            idx = i
            break
    nearby = "\n".join(lines[idx: idx + 18]) if lines else text

    name = first_match([
        r"(?:Amministratore\s+unico|Legale\s+rappresentante|Titolare|Presidente)\s*[:\-]?\s*([A-ZÀ-Ü' ]{5,80})",
        r"Cognome\s+e\s+nome\s*[:\-]?\s*([A-ZÀ-Ü' ]{5,80})",
        r"Nome\s*[:\-]?\s*([A-ZÀ-Ü' ]{5,80})",
    ], nearby, flags=re.I | re.M)
    name = re.sub(r"\s+(NATO|NATA|CODICE|CARICA|RESIDENTE).*$", "", name, flags=re.I).strip()

    if not name:
        for line in nearby.splitlines():
            if re.fullmatch(r"[A-ZÀ-Ü' ]{5,80}", line.strip()) and not re.search(r"AMMINISTRATORE|CARICA|NATO|CODICE|FISCALE", line, re.I):
                name = clean_inline(line)
                break

    carica = first_match([
        r"Carica\s*[:\-]?\s*([^\n]{3,80})",
        r"(Amministratore\s+unico|Presidente\s+consiglio\s+amministrazione|Consigliere\s+delegato|Legale\s+rappresentante)",
    ], nearby)
    if not carica and re.search(r"amministratore unico", nearby, re.I):
        carica = "Amministratore unico"

    cf = first_match([
        r"Codice\s+fiscale\s*[:\-]?\s*([A-Z0-9]{16})",
        r"C\.?F\.?\s*[:\-]?\s*([A-Z0-9]{16})",
    ], nearby)
    if not cf:
        m = re.search(r"\b([A-Z]{6}[0-9]{2}[A-Z][0-9]{2}[A-Z][0-9]{3}[A-Z])\b", admin_block, re.I)
        cf = m.group(1).upper() if m else ""

    birth = first_match([
        r"nat[oa]\s+a\s+([^\n,;]+?)\s+il\s+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
    ], nearby)
    luogo = ""
    data_n = ""
    m_birth = re.search(r"nat[oa]\s+a\s+([^\n,;]+?)\s+il\s+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", nearby, re.I)
    if m_birth:
        luogo = clean_inline(m_birth.group(1))
        data_n = clean_inline(m_birth.group(2))
    else:
        luogo = first_match([r"Luogo\s+di\s+nascita\s*[:\-]?\s*([^\n]{2,80})"], nearby)
        data_n = first_match([r"Data\s+di\s+nascita\s*[:\-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})"], nearby)

    residenza = first_match([
        r"residente\s+(?:in|a)\s+([^\n;]{5,160})",
        r"domiciliat[oa]\s+(?:in|a)\s+([^\n;]{5,160})",
        r"Residenza\s*[:\-]?\s*([^\n]{5,160})",
        r"Domicilio\s*[:\-]?\s*([^\n]{5,160})",
    ], nearby)

    return {
        "amministratore_nome": name.title() if name.isupper() else name,
        "amministratore_carica": carica,
        "amministratore_codice_fiscale": cf.upper(),
        "amministratore_luogo_nascita": luogo,
        "amministratore_data_nascita": data_n,
        "amministratore_residenza": residenza,
    }


def parse_document(pdf_bytes: bytes) -> Dict[str, str]:
    text = extract_pdf_text_from_bytes(pdf_bytes)
    lines = lines_from_text(text)
    mail, pec = extract_emails(text)
    piva = extract_piva(text)
    ateco, settore = extract_ateco(lines, text)
    data: Dict[str, str] = {
        "denominazione": extract_denominazione(lines, text),
        "partita_iva": piva,
        "codice_fiscale": extract_company_cf(text, piva),
        "sede_legale": extract_sede_legale(lines, text),
        "pec": pec,
        "mail_azienda": mail,
        "codice_ateco": ateco,
        "settore_attivita": settore,
    }
    data.update(extract_admin(lines, text))
    data["testo_estratto"] = text[:8000]
    return data


# -----------------------------------------------------------------------------
# Compilazione DOCX
# -----------------------------------------------------------------------------


def piva_cf_text(client: Dict[str, str]) -> str:
    piva = re.sub(r"\D+", "", client.get("partita_iva", ""))
    cf = clean_inline(client.get("codice_fiscale", "")).upper()
    sede = clean_inline(client.get("sede_legale", ""))
    pec = clean_inline(client.get("pec", ""))

    if piva and cf and piva != cf:
        base = f"{piva} - C.F. {cf}"
    else:
        base = piva or cf

    extra = []
    if sede:
        extra.append(f"sede legale {sede}")
    if pec:
        extra.append(f"PEC {pec}")

    if base and extra:
        return base + ", " + ", ".join(extra)
    return base or ", ".join(extra)


def build_admin_intro(client: Dict[str, str]) -> str:
    parts: List[str] = []
    if client.get("amministratore_nome"):
        parts.append(client["amministratore_nome"])
    luogo = client.get("amministratore_luogo_nascita", "")
    data_n = client.get("amministratore_data_nascita", "")
    if luogo and data_n:
        parts.append(f"nato a {luogo} il {data_n}")
    elif luogo:
        parts.append(f"nato a {luogo}")
    elif data_n:
        parts.append(f"nato il {data_n}")
    if client.get("amministratore_residenza"):
        parts.append(f"residente/domiciliato in {client['amministratore_residenza']}")
    if client.get("amministratore_codice_fiscale"):
        parts.append(f"codice fiscale {client['amministratore_codice_fiscale']}")
    return "; ".join([p for p in parts if p])


def build_descrizione_compenso(importo_finanziamento: str, banca: str) -> str:
    imp = clean_inline(importo_finanziamento)
    banca = clean_inline(banca)
    if imp and banca:
        return f"consulenza e assistenza per richiesta, istruttoria e ottenimento di finanziamento pari a {imp} presso {banca}"
    if imp:
        return f"consulenza e assistenza per richiesta, istruttoria e ottenimento di finanziamento pari a {imp}"
    if banca:
        return f"consulenza e assistenza per richiesta, istruttoria e ottenimento di finanziamento presso {banca}"
    return "consulenza e assistenza finanziaria, analisi documentale e predisposizione fascicolo bancario"


def replace_text_in_paragraph(paragraph: Any, replacements: Dict[str, str]) -> None:
    if not paragraph.runs:
        return
    full = "".join(run.text for run in paragraph.runs)
    new = full
    for src, dst in replacements.items():
        if src in new:
            new = new.replace(src, dst)
    # Normalizzazioni mirate per segnaposti con spazi o apostrofi diversi
    new = re.sub(r"INSERIRE\s+DATI\s+AMINISTRATORE", replacements.get("INSERIRE DATI AMINISTRATORE", ""), new, flags=re.I)
    new = re.sub(r"INSERIRE\s+DATI\s+AMMINISTRATORE", replacements.get("INSERIRE DATI AMINISTRATORE", ""), new, flags=re.I)
    new = re.sub(r"INSERIRE\s+DATI\s+SOCIETA[’'`]", replacements.get("INSERIRE DATI SOCIETA’", ""), new, flags=re.I)
    new = re.sub(r"INSERIRE\s+PARTITA\s+IVA", replacements.get("INSERIRE PARTITA IVA", ""), new, flags=re.I)
    new = re.sub(r"INSERIRE\s+LA\s+DATA", replacements.get("INSERIRE LA DATA", today_it()), new, flags=re.I)
    if new != full:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ""


def iter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def compile_docx(template_bytes: bytes, client: Dict[str, str], mandate: Dict[str, str]) -> bytes:
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        template_path = td_path / "template.docx"
        output_path = td_path / "mandato_compilato.docx"
        template_path.write_bytes(template_bytes)
        doc = Document(str(template_path))

        importo_consulenza = format_euro(mandate.get("importo_consulenza", ""))
        importo_finanziamento = format_euro(mandate.get("importo_finanziamento", ""))
        descrizione = clean_inline(mandate.get("descrizione_compenso", "")) or build_descrizione_compenso(importo_finanziamento, mandate.get("banca", ""))

        replacements = {
            "INSERIRE DATI AMINISTRATORE": build_admin_intro(client),
            "INSERIRE DATI AMMINISTRATORE": build_admin_intro(client),
            "INSERIRE DATI SOCIETA’": clean_inline(client.get("denominazione", "")),
            "INSERIRE DATI SOCIETA'": clean_inline(client.get("denominazione", "")),
            "INSERIRE PARTITA IVA": piva_cf_text(client),
            "INSERIRE LA DATA": clean_inline(mandate.get("data_mandato", "")) or today_it(),
            "PARI ad  importo": f"PARI ad {importo_consulenza}" if importo_consulenza else "PARI ad",
            "PARI ad importo": f"PARI ad {importo_consulenza}" if importo_consulenza else "PARI ad",
            "PARI AD IMPORTO": f"PARI ad {importo_consulenza}" if importo_consulenza else "PARI ad",
            "IMPORTO": importo_consulenza,
            "FINANZIAMENTO": importo_finanziamento,
            "BANCA": clean_inline(mandate.get("banca", "")),
            "INSERIRE DESCRIZIONE": descrizione,
        }

        for paragraph in iter_paragraphs(doc):
            replace_text_in_paragraph(paragraph, replacements)

        doc.save(str(output_path))
        return output_path.read_bytes()


def convert_docx_bytes_to_pdf(docx_bytes: bytes) -> Optional[bytes]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        source = td_path / "mandato.docx"
        source.write_bytes(docx_bytes)
        cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(td_path), str(source)]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=90)
            pdf_path = td_path / "mandato.pdf"
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                return pdf_path.read_bytes()
        except Exception:
            return None
    return None


# -----------------------------------------------------------------------------
# UI Streamlit
# -----------------------------------------------------------------------------


def load_default_template() -> bytes:
    if DEFAULT_TEMPLATE.exists():
        return DEFAULT_TEMPLATE.read_bytes()
    raise FileNotFoundError("Manca assets/Mandato_vuoto.docx")


def render_field(label: str, key: str, data: Dict[str, str], col=None) -> str:
    target = col if col is not None else st
    return target.text_input(label, value=clean_inline(data.get(key, "")), key=f"field_{key}")


def main() -> None:
    st.set_page_config(page_title=APP_NAME, page_icon="📄", layout="wide")

    st.markdown(
        """
        <style>
        .block-container {padding-top: 1.2rem;}
        .fp-title {font-size: 2.0rem; font-weight: 800; color: #0b2341; margin-bottom: 0.1rem;}
        .fp-subtitle {font-size: 1rem; color: #8b5a2b; margin-bottom: 1rem;}
        .stButton>button, .stDownloadButton>button {border-radius: 8px; font-weight: 700;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    top_left, top_right = st.columns([6, 1])
    with top_left:
        st.markdown('<div class="fp-title">FinancePlus Mandato Web</div>', unsafe_allow_html=True)
        st.markdown('<div class="fp-subtitle">Versione GitHub / Streamlit Cloud senza tkinter</div>', unsafe_allow_html=True)
    with top_right:
        if LOGO_FILE.exists():
            st.image(str(LOGO_FILE), use_container_width=True)

    st.info(
        "Questa versione risolve l'errore Streamlit Cloud: il file desktop Tkinter non va eseguito sul cloud. "
        "Qui il main file corretto e' `streamlit_app.py` oppure il wrapper `Mandato_aggiornato_report.py`."
    )

    if "client_data" not in st.session_state:
        st.session_state.client_data = {}

    tab1, tab2, tab3 = st.tabs(["1. Inserisci report/visura PDF", "2. Anagrafica e mandato", "3. Download e guida"])

    with tab1:
        st.subheader("Carica PDF")
        uploaded_pdf = st.file_uploader("Visura camerale o report PDF", type=["pdf"], key="uploaded_pdf")
        if uploaded_pdf is not None:
            pdf_bytes = uploaded_pdf.getvalue()
            with st.spinner("Lettura PDF ed estrazione dati..."):
                parsed = parse_document(pdf_bytes)
            st.session_state.client_data.update({k: v for k, v in parsed.items() if k != "testo_estratto" and v})
            st.success("Dati estratti. Controllali nella scheda Anagrafica e mandato.")
            with st.expander("Anteprima testo estratto"):
                st.text_area("Testo PDF", value=parsed.get("testo_estratto", ""), height=260)

        st.markdown("### Importa/esporta anagrafica")
        uploaded_json = st.file_uploader("Carica anagrafica JSON salvata in precedenza", type=["json"], key="uploaded_json")
        if uploaded_json is not None:
            try:
                st.session_state.client_data.update(json.loads(uploaded_json.getvalue().decode("utf-8")))
                st.success("Anagrafica JSON caricata.")
            except Exception as exc:
                st.error(f"JSON non valido: {exc}")

    with tab2:
        st.subheader("Dati azienda")
        data = dict(st.session_state.client_data)
        c1, c2 = st.columns(2)
        data["denominazione"] = render_field("Denominazione società", "denominazione", data, c1)
        data["partita_iva"] = render_field("Partita IVA", "partita_iva", data, c2)
        data["codice_fiscale"] = render_field("Codice fiscale società", "codice_fiscale", data, c1)
        data["pec"] = render_field("PEC", "pec", data, c2)
        data["sede_legale"] = render_field("Sede legale", "sede_legale", data)
        c3, c4 = st.columns(2)
        data["mail_azienda"] = render_field("Email ordinaria", "mail_azienda", data, c3)
        data["codice_ateco"] = render_field("Codice ATECO", "codice_ateco", data, c4)
        data["settore_attivita"] = render_field("Descrizione attività", "settore_attivita", data)

        st.subheader("Dati amministratore")
        c5, c6 = st.columns(2)
        data["amministratore_nome"] = render_field("Nome amministratore", "amministratore_nome", data, c5)
        data["amministratore_carica"] = render_field("Carica", "amministratore_carica", data, c6)
        data["amministratore_codice_fiscale"] = render_field("Codice fiscale amministratore", "amministratore_codice_fiscale", data, c5)
        data["amministratore_luogo_nascita"] = render_field("Luogo nascita", "amministratore_luogo_nascita", data, c6)
        data["amministratore_data_nascita"] = render_field("Data nascita", "amministratore_data_nascita", data, c5)
        data["amministratore_residenza"] = render_field("Residenza/domicilio", "amministratore_residenza", data, c6)

        st.session_state.client_data = data

        st.subheader("Dati mandato")
        m1, m2, m3 = st.columns(3)
        mandate = {
            "data_mandato": m1.text_input("Data mandato", value=today_it()),
            "importo_consulenza": m2.text_input("Compenso consulenza", value=""),
            "importo_finanziamento": m3.text_input("Importo finanziamento", value=""),
            "banca": st.text_input("Banca / intermediario", value=""),
            "descrizione_compenso": st.text_area("Descrizione compenso / incarico", value="", height=90),
        }

        uploaded_template = st.file_uploader("Modello Word personalizzato opzionale (.docx)", type=["docx"], key="uploaded_template")
        template_bytes = uploaded_template.getvalue() if uploaded_template is not None else load_default_template()

        if st.button("Genera mandato DOCX", type="primary"):
            if not data.get("denominazione"):
                st.error("Inserisci almeno la denominazione della società.")
            else:
                try:
                    docx_bytes = compile_docx(template_bytes, data, mandate)
                    st.session_state.generated_docx = docx_bytes
                    st.session_state.generated_name = safe_filename(f"Mandato_{data.get('denominazione', 'cliente')}")
                    st.success("Mandato DOCX generato correttamente.")
                except Exception as exc:
                    st.error(f"Errore generazione DOCX: {exc}")

    with tab3:
        st.subheader("Download")
        name = st.session_state.get("generated_name", "Mandato_cliente")
        docx_bytes = st.session_state.get("generated_docx")
        if docx_bytes:
            st.download_button(
                "Scarica mandato DOCX",
                data=docx_bytes,
                file_name=f"{name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
            )
            pdf_bytes = convert_docx_bytes_to_pdf(docx_bytes)
            if pdf_bytes:
                st.download_button("Scarica mandato PDF", data=pdf_bytes, file_name=f"{name}.pdf", mime="application/pdf")
            else:
                st.warning("PDF non disponibile: su Cloud serve LibreOffice. Il DOCX è comunque generato.")
        else:
            st.write("Genera prima il mandato nella scheda 2.")

        client_json = json.dumps(st.session_state.client_data, ensure_ascii=False, indent=2).encode("utf-8")
        st.download_button("Scarica anagrafica JSON", data=client_json, file_name="anagrafica_cliente.json", mime="application/json")

        st.markdown(
            """
            ### Come pubblicare su Streamlit Cloud
            1. Carica tutto il contenuto della cartella su GitHub.
            2. Su Streamlit Cloud scegli il repository.
            3. Come **Main file path** indica `streamlit_app.py`.
            4. In alternativa puoi indicare `Mandato_aggiornato_report.py`, perché ora è solo un wrapper web e non importa tkinter.
            5. Non usare il file desktop dentro la cartella `desktop/` su Streamlit Cloud.
            """
        )


if __name__ == "__main__":
    main()
